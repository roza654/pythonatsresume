from flask import (
    Flask, render_template, request, jsonify,
    session, redirect, url_for, flash
)
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash

from docx import Document
import pdfplumber
import re
import io
import os

from ats_engine import analyze_resume
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# ================= FLASK APP =================
app = Flask(
    __name__,
    template_folder="frontend/templates",
    static_folder="frontend/static"
)

# ================= CONFIG =================
app.config['SECRET_KEY'] = 'secret123'
app.config['SQLALCHEMY_DATABASE_URI'] = (
    "mysql+pymysql://ats_user:roza123@localhost/ats_db"

)
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# ================= CONSTANTS =================
MAX_FILE_SIZE = 5 * 1024 * 1024
ALLOWED_EXTS = (".pdf", ".docx")

# ================= DATABASE MODELS =================
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)

# ================= ROUTES =================
@app.route("/test-db")
def test_db():
    return "Flask + SQLAlchemy Connected 🚀"

@app.route("/")
def home():
    return render_template("home.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")

        user = User.query.filter_by(email=email).first()

        if user and check_password_hash(user.password, password):
            session["user_id"] = user.id
            session["user_email"] = user.email
            return redirect(url_for("dashboard"))
        else:
            flash("Invalid email or password")

    return render_template("login.html")

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")

        if User.query.filter_by(email=email).first():
            flash("User already exists")
            return redirect(url_for("signup"))

        hashed = generate_password_hash(password)
        user = User(email=email, password=hashed)
        db.session.add(user)
        db.session.commit()

        flash("Account created successfully")
        return redirect(url_for("login"))

    return render_template("signup.html")

@app.route("/dashboard")
def dashboard():
    if "user_id" not in session:
        return redirect(url_for("login"))
    return render_template("dashboard.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

@app.route("/build-resume")
def build_resume():
    return render_template("index.html")

@app.route("/ats-result")
def ats_result():
    return render_template("atsanalysis.html")

# ================= RESUME EXTRACTION (❌ TOUCH NAHI KIYA) =================
@app.route("/extract_resume", methods=["POST"])
def extract_resume():
    try:
        file = request.files.get("resume")
        if not file:
            return jsonify({"error": "no file uploaded"}), 400

        filename = file.filename.lower()
        if not filename.endswith(ALLOWED_EXTS):
            return jsonify({"error": "unsupported file type"}), 400

        file_bytes = file.read()
        if len(file_bytes) > MAX_FILE_SIZE:
            return jsonify({"error": "file too large"}), 413

        text = ""

        if filename.endswith(".docx"):
            doc = Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())
            text = "\n".join(parts)

        elif filename.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = [p.extract_text() for p in pdf.pages if p.extract_text()]
                text = "\n".join(pages)

        lines = [l.strip() for l in text.split("\n") if l.strip()]

        email_match = re.search(
            r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text
        )
        email = email_match.group(0) if email_match else ""

        phone_match = re.search(r"(\+\d{1,3}[-\s]?)?\d{7,15}", text)
        phone = phone_match.group(0) if phone_match else ""

        return jsonify({
            "email": email,
            "phone": phone,
            "summary": text[:900]
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 400

# ================= ATS ANALYSIS (❌ TOUCH NAHI KIYA) =================
@app.route("/analyze", methods=["POST"])
def analyze():
    data = request.json or {}
    resume = data.get("resume", "")
    jd = data.get("job_description", "")

    if not resume or not jd:
        return jsonify({"overall_score": 0})

    vectorizer = TfidfVectorizer(stop_words="english")
    vectors = vectorizer.fit_transform([resume, jd])
    similarity = cosine_similarity(vectors[0], vectors[1])[0][0]

    score = round(similarity * 100, 2)
    return jsonify({"overall_score": score})

# ================= RUN (SAFE) =================
if __name__ == "__main__":
    with app.app_context():
        db.create_all()
        print("✅ Tables ensured (one time)")
    app.run(debug=False)
